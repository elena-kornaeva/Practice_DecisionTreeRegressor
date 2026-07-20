from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from pathlib import Path
p=Path(r'C:\Users\bulki\Downloads\Telegram Desktop\титульник.docx')
d=Document(p)
s=d.sections[0]
print('PAGE',s.page_width,s.page_height,'MARGINS',s.top_margin,s.right_margin,s.bottom_margin,s.left_margin)
for i,p in enumerate(d.paragraphs):
    pf=p.paragraph_format
    print(f'P{i} text={p.text!r} align={p.alignment} before={pf.space_before} after={pf.space_after} line={pf.line_spacing} left={pf.left_indent} right={pf.right_indent} first={pf.first_line_indent} keep_next={pf.keep_with_next} page_break={pf.page_break_before}')
    for j,r in enumerate(p.runs):
        print('  R',j,repr(r.text),'font',r.font.name,'size',r.font.size,'bold',r.bold,'italic',r.italic,'all_caps',r.font.all_caps,'underline',r.underline)
for ti,t in enumerate(d.tables):
    print('TABLE',ti,'alignment',t.alignment,'autofit',t.autofit,'rows',len(t.rows),'cols',len(t.columns))
    tblPr=t._tbl.tblPr
    print(' tblW',tblPr.tblW.type,tblPr.tblW.w)
    for ri,row in enumerate(t.rows):
        print(' row',ri,'height',row.height,'rule',row.height_rule)
        for ci,c in enumerate(row.cells):
            tcW=c._tc.tcPr.tcW
            print('  cell',ci,'width',tcW.type,tcW.w,'text',repr(c.text))
            for p in c.paragraphs:
                print('   P align',p.alignment,'before',p.paragraph_format.space_before,'after',p.paragraph_format.space_after)
                for r in p.runs: print('    R',repr(r.text),r.font.name,r.font.size,r.bold,r.italic)
